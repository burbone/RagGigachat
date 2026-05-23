from __future__ import annotations

import logging
import re
from pathlib import Path

from src.models import DocType, PageBlock, PageType, DocumentStats

logger = logging.getLogger(__name__)

_MIN_TEXT_LEN = 50
TABLE_MARKER_RE = re.compile(r"\|.+\|")


def _clean_footnote_artifacts(text):
    text = re.sub(r'(\d+)[1-9]\d?(?=\s|руб|$)', lambda m: m.group(1), text)
    text = re.sub(r'([а-яёА-ЯЁa-zA-Z])\d{1,2}(?=\s|[.,]|$)', r'\1', text)
    return text


def _rows_to_markdown(rows):
    if not rows:
        return ""
    cleaned = [[str(c or "").strip() for c in row] for row in rows]
    cleaned = [r for r in cleaned if any(r)]
    if not cleaned:
        return ""
    header = cleaned[0]
    ncols = len(header)
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * ncols) + " |",
    ]
    for row in cleaned[1:]:
        row = (row + [""] * ncols)[:ncols]
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _load_pdf(path):
    import pymupdf4llm
    import pdfplumber

    logger.info(f"PDF {path.name}: первый проход через pymupdf4llm")
    raw = pymupdf4llm.to_markdown(str(path), page_chunks=True)

    pages = []
    skipped = 0

    with pdfplumber.open(str(path)) as pdf_doc:
        for chunk in raw:
            page_num = chunk["metadata"].get("page_number", 0)
            text = chunk["text"].strip().replace("<br>", " ").replace("<br/>", " ")
            has_tbl = bool(TABLE_MARKER_RE.search(text))

            if len(text) >= _MIN_TEXT_LEN:
                pages.append(PageBlock(
                    page=page_num, text=text,
                    source=path.name, has_tables=has_tbl,
                ))
                continue

            idx = max(0, page_num - 1)
            if idx < len(pdf_doc.pages):
                pl = pdf_doc.pages[idx]
                try:
                    pl_text = (pl.extract_text() or "").strip()
                    tables = pl.extract_tables() or []
                    tbl_md = "\n\n".join(_rows_to_markdown(t) for t in tables if t)
                    combined = "\n\n".join(filter(None, [pl_text, tbl_md]))
                    if len(combined) >= _MIN_TEXT_LEN:
                        pages.append(PageBlock(
                            page=page_num, text=combined,
                            source=path.name, has_tables=bool(tables),
                        ))
                        continue
                except Exception as e:
                    logger.debug(f"pdfplumber стр.{page_num}: {e}")

            ocr_text = _ocr_page(str(path), page_num)
            if ocr_text and len(ocr_text) >= _MIN_TEXT_LEN:
                pages.append(PageBlock(
                    page=page_num, text=ocr_text, source=path.name,
                ))
                continue

            skipped += 1
            logger.debug(f"Стр.{page_num} пропущена")

    logger.info(
        f"PDF {path.name} → страниц={len(pages)}  с_таблицами={sum(p.has_tables for p in pages)}  пропущено={skipped}"
    )
    return pages


def _ocr_page(pdf_path, page_num):
    try:
        import fitz
        import pytesseract
        from PIL import Image
        import io

        doc = fitz.open(pdf_path)
        page = doc[max(0, page_num - 1)]
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        doc.close()
        result = pytesseract.image_to_string(img, lang="rus+eng")
        logger.debug(f"OCR стр.{page_num}: {len(result)} символов")
        return result.strip() or None
    except ImportError:
        return None
    except Exception as e:
        logger.warning(f"OCR стр.{page_num} ошибка: {e}")
        return None


def _load_docx(path):
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Установите: pip install python-docx")

    doc = Document(str(path))
    parts = []
    has_tbl = False

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        has_tbl = True
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        md = _rows_to_markdown(rows)
        if md:
            parts.append(md)

    full = "\n\n".join(parts)
    if not full.strip():
        return []

    # DOCX не имеет страниц
    pages = []
    size = 3000
    for i, start in enumerate(range(0, len(full), size)):
        pages.append(PageBlock(
            page=i + 1,
            text=full[start:start + size],
            source=path.name,
            doc_type=DocType.DOCX,
            has_tables=has_tbl,
        ))

    logger.info(f"DOCX {path.name} → блоков={len(pages)}")
    return pages


def _load_text(path):
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        return []

    dtype = DocType.MD if path.suffix.lower() in (".md", ".markdown") else DocType.TXT
    pages = []
    size = 3000
    for i, start in enumerate(range(0, len(text), size)):
        pages.append(PageBlock(
            page=i + 1,
            text=text[start:start + size],
            source=path.name,
            doc_type=dtype,
        ))

    logger.info(f"TXT/MD {path.name} → блоков={len(pages)}")
    return pages


_LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".doc": _load_docx,
    ".txt": _load_text,
    ".md": _load_text,
    ".markdown": _load_text,
}

SUPPORTED_EXTENSIONS = set(_LOADERS.keys())


class DocumentLoader:
    def load(self, file_path):
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        ext = path.suffix.lower()
        fn = _LOADERS.get(ext)
        if fn is None:
            raise ValueError(
                f"Неподдерживаемый формат '{ext}'. "
                f"Поддерживаются: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        logger.info(f"Загружаю {path.name} ...")
        pages = fn(path)

        table_pages = sum(1 for p in pages if p.has_tables)
        stats = DocumentStats(
            total_pages=len(pages),
            parsed_pages=len(pages),
            text_pages=len(pages) - table_pages,
            table_pages=table_pages,
            total_tables=table_pages,
            image_only_pages=0,
        )
        return pages, stats